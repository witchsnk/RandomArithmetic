#!/usr/bin/python
# -*- coding: UTF-8 -*-

import math
import random
from docx import Document
# from docx.shared import Pt
from stack import Stack


## 输入条件
def inputCondition():
    print("输入生成的种类，混合的用连续数字组合即可（1:加法；2:减法；3:乘法；4:除法）")
    ruleType = input("请输入：") or "1"

    print("输入计算式的混合数量（2～4）")
    digit = int(input("请输入：") or 2)

    print("输入计算式数值范围（1～100）")
    scope = int(input("请输入：") or 1)

    print("是否允许负数出现(1:允许 0:不允许<默认>)")
    neg = int(input("请输入：") or 0)

    print("生成试题的数量（默认120道）")
    num = int(input("请输入：") or 120)

    print("用户输入的生成规则需要生成 ")
    return ruleType, digit, scope, neg, num

def analyseRule():
    pass

def randomArithmetic(ruleType, digit, scope, neg, num):
    data = []
    for i in range(0, num):
        jump = True
        while(jump):
            result = madeArithmetic(ruleType, digit, scope)
            if result[0] == True:
                dataSubList = result[1]
                jump = False
        # 转换成str输出
        dataStr = "".join(dataSubList)
        dataStr = dataStr + str(calcResult(getPostfix(dataSubList[:-1])))
        data.append(dataStr)
    return data
## 产生计算式
def madeArithmetic(ruleType, digit, scope):
    # 先随机计算式参与计算的数有几个
    n = random.randint(2, digit)
    # 创建一个生成数2倍的数组
    dataSubList = [0] * (n * 2)
    # 隔位把数插入单数位置
    for x in range(0, n * 2, 2):
        dataSubList[x] = str(random.randint(0, scope))
    for x in range(1, n * 2 - 1, 2):
        ruleList = list(ruleType)
        d = int(random.sample(ruleList, 1)[0])
        if d == 1:
            dataSubList[x] = "+"
        elif d == 2:
            dataSubList[x] = "-"
        elif d == 3:
            dataSubList[x] = "*"
        elif d == 4:
            dataSubList[x] = "/"
    
    # 判断是不是有 * 和 / ，/的有面不能是0

    # 如果不能出现负数 就要判断“-”前面的数一定要大于右面的数
    # 有几种情况
    # a - b         a>b
    # a - b - c     a>b+c
    # a - b - c - d a>b+c+d
    # a + b - c     a+b>c   
    # if neg == 0:
    #     if dataSubList.count("-") > 0:
    #         while(True):
    #             pass
            #     if 
            # i = dataSubList.index("-")
            # temp = dataSubList[i - 1]
            # dataSubList[i - 1] = dataSubList[i + 1]
            # dataSubList[i + 1] = temp
    # 隔位插入运算符号 最后补上“=”
    dataSubList[n * 2 - 1] = "="
    return True, dataSubList

## 获取运算符的优先级
def prior(c):
    if c == "+" or c == "-":
        return 1
    elif c == "*" or c == "/":
        return 2
    else:
        return 0

## 判断是否是运算符
def isOperator(c):
    if c == "+" or c == "-" or c == "*" or c == "/":
        return True
    else:
        return False
## 中缀转后缀
def getPostfix(expr):
    s = Stack()
    output = []
    for c in expr:
        if isOperator(c):
            while(not s.empty() and isOperator(s.top()) and prior(s.top()) >= prior(c)):
                output.append(s.top())
                s.pop()
            s.push(c)
        elif c == "(":
            s.push(c)
        elif c == ")":
            while s.top() != "(":
                output.append(s.top())
                s.pop()
            s.pop()
        else:
            output.append(c)
    while not s.empty():
        output.append(s.top())
        s.pop()
    # print(output)
    return output
## 从栈中连续弹出两个操作数
def popTwoNumbers(stack):
    aa = stack.top()
    stack.pop()
    bb = stack.top()
    stack.pop()
    return aa, bb

 ## 计算出结果
def calcResult(expr):
    first = 0
    second = 0
    s = Stack()
    for c in expr:
        if c == "+":
            first,second = popTwoNumbers(s)
            s.push(int(second) + int(first))
            # print("{0} + {1}",second,first)
        elif c == "-":
            first,second = popTwoNumbers(s)
            s.push(int(second) - int(first))
            # print("{0} - {1}",second,first)
        elif c == "*":
            first,second = popTwoNumbers(s)
            s.push(int(second) * int(first))
            # print("{0} * {1}",second,first)
        elif c == "/":
            first,second = popTwoNumbers(s)
            s.push(int(second) / int(first))
            # print("{0} / {1}",second,first)
        else:
            s.push(c)
    result = s.top()
    s.pop
    return result

## 输出到word
def outPutWord(data):
    document = Document()
    document.add_heading('小学生数学计算练习题')
    tab = document.add_table(rows=23, cols=4)
    item = 0
    for x in range(0,23):
        for y in range(0,4):
            cell = tab.cell(x,y)
            cell.text = data[item]
            item += 1
    document.save('test.docx')


if __name__ == "__main__":
    ruleType, digit, scope, neg, num = inputCondition()
    data = randomArithmetic(ruleType, digit, scope, neg, num)
    outPutWord(data)

    # calcStr = "1+2*5+(3-1)"
    # print(calcStr)
    # s = getPostfix(calcStr)
    # print(s)
    # r = calcResult(s)
    # print(str(r))
    pass